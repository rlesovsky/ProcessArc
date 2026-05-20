"""Treating Sequence Word document exporter (redesigned — Tier 1).

Produces `{SiteName}_TreatingSequence.docx` from the customer's sequence
workbook — the same prose the engineer confirmed at Discover / Extract,
rendered as a structured, readable Word document.

The customer's sequence sheets are running English. Earlier this exporter
emitted that prose as a flat list of paragraphs with a crude bold heuristic.
This version classifies every line (see `sequence_classify.py`) and renders
each kind distinctively:

  - Cylinder sections become numbered steps. Each step's actions are a
    numbered list; the transition condition that *ends* the step is rendered
    as a separate "ADVANCES WHEN" line (previously it was wrongly bolded as if
    it were a step header).
  - Conditional actions ("If treating Tank 1, open S3") nest under the action
    they qualify.
  - Mix sections, which the customer writes as a flat numbered procedure
    rather than named phases, render as one numbered procedure with any
    "For mixing ..." sub-group labels as subheadings.
  - Customer notes and open questions ("confirm in PLC?", simultaneous-mix
    requests) are pulled into shaded callouts so they stand out.
  - A cover page, table of contents, and per-section metadata strip make the
    document a shareable deliverable.

Text fidelity: the customer's wording is reproduced verbatim except for a
small, closed set of unambiguous spelling corrections, each marked in the text
and listed in an Editorial Notes section. Device names, numbers, and
operational content are never altered. The cover carries a standing note to
this effect — important because this document reads the raw workbook, so it
reflects the source sheet, not any post-review device-model corrections.

Public interface is unchanged: `export_sequence_doc(plant, workbook_path,
output_dir)` returns the output `Path`, or `None` if the source workbook is
missing — callers handle that as a missing-input case, not a hard failure.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.export.sequence_classify import (
	ClassifiedLine,
	LineType,
	classify_lines,
)
from backend.ingest import IngestedSheet, SheetKind, ingest_workbook
from backend.model import DeviceClass, DeviceModel, ReviewStatus
from backend.model.device import DeviceRecord, SystemKind
from backend.model.plant import PlantConfiguration


# --- palette -----------------------------------------------------------------
_INK = RGBColor(0x1A, 0x1A, 0x1A)
_NAVY = RGBColor(0x1F, 0x3A, 0x5F)
_STEEL = RGBColor(0x2E, 0x75, 0xB6)
_GREY = RGBColor(0x70, 0x70, 0x70)
_AMBER_TEXT = RGBColor(0x5A, 0x42, 0x10)
_AMBER_LABEL = RGBColor(0x9A, 0x6A, 0x12)
_COND_LABEL = RGBColor(0xB0, 0x7A, 0x1E)
_TRANS_TEXT = RGBColor(0x44, 0x44, 0x44)

_AMBER_BG = "FBE7C8"
_NOTE_BG = "EAF1F7"
_STEP_BG = "EDEDED"
_RULE = "C9C9C9"
_STEEL_HEX = "2E75B6"


# --- low-level helpers -------------------------------------------------------
def _strip_invalid_zoom(doc: Document) -> None:
	"""Remove the schema-incomplete <w:zoom> element python-docx inserts.

	python-docx (1.2.0) writes a default `<w:zoom w:val="bestFit"/>` into
	settings.xml with no `w:percent` attribute. Word opens this fine, but a
	strict OOXML validator rejects it. Dropping the element entirely is safe —
	Word falls back to its own default zoom — and lets the exported file pass
	validation cleanly.
	"""
	settings = doc.settings.element
	for zoom in settings.findall(qn("w:zoom")):
		settings.remove(zoom)


def _shade(paragraph, hex_fill: str) -> None:
	"""Apply a solid background shade to a paragraph."""
	pPr = paragraph._p.get_or_add_pPr()
	shd = OxmlElement("w:shd")
	shd.set(qn("w:val"), "clear")
	shd.set(qn("w:color"), "auto")
	shd.set(qn("w:fill"), hex_fill)
	pPr.append(shd)


def _border(paragraph, color: str = _RULE, size: int = 6) -> None:
	"""Draw a bottom rule under a paragraph (used as section dividers)."""
	pPr = paragraph._p.get_or_add_pPr()
	bdr = OxmlElement("w:pBdr")
	bottom = OxmlElement("w:bottom")
	bottom.set(qn("w:val"), "single")
	bottom.set(qn("w:sz"), str(size))
	bottom.set(qn("w:space"), "4")
	bottom.set(qn("w:color"), color)
	bdr.append(bottom)
	pPr.append(bdr)


def _run(
	paragraph,
	text: str,
	*,
	bold: bool = False,
	italic: bool = False,
	size: float = 10.5,
	color: RGBColor = _INK,
	caps: bool = False,
):
	r = paragraph.add_run(text)
	r.bold = bold
	r.italic = italic
	r.font.size = Pt(size)
	r.font.color.rgb = color
	if caps:
		r.font.all_caps = True
	return r


def _mark_correction(paragraph, line: ClassifiedLine) -> None:
	"""Append the small correction mark if this line had a typo fixed."""
	if line.corrections:
		_run(paragraph, "  \u2731", size=7, color=_STEEL)


def _safe_site_name(plant: PlantConfiguration) -> str:
	base = plant.site_name.strip() or "Project"
	return "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in base)


# --- section heading derivation (unchanged behavior) -------------------------
def _section_heading(sheet: IngestedSheet) -> str:
	if sheet.kind == SheetKind.CYLINDER_SEQUENCE and sheet.cylinder_number is not None:
		return f"Cylinder {sheet.cylinder_number}"
	if sheet.kind == SheetKind.MIX_SEQUENCE:
		label = (sheet.mix_label or "").strip()
		return f"{label} Mix" if label else "Mix"
	return sheet.name


def _cylinder_is_idle(plant: PlantConfiguration, sheet: IngestedSheet) -> bool:
	"""True if this cylinder is flagged idle in the Plant Configuration."""
	if sheet.kind != SheetKind.CYLINDER_SEQUENCE or sheet.cylinder_number is None:
		return False
	for cyl in plant.cylinders:
		if cyl.number == sheet.cylinder_number:
			return cyl.is_idle
	return False


# --- document pieces ---------------------------------------------------------
def _setup_styles(doc: Document) -> None:
	normal = doc.styles["Normal"]
	normal.font.name = "Calibri"
	normal.font.size = Pt(10.5)
	normal.font.color.rgb = _INK
	_strip_invalid_zoom(doc)


def _add_cover(doc: Document, plant: PlantConfiguration) -> None:
	site = plant.site_name or "Project"
	for _ in range(3):
		doc.add_paragraph()

	p = doc.add_paragraph()
	_run(p, "TREATING SEQUENCE", bold=True, size=13, color=_STEEL)
	p.paragraph_format.space_after = Pt(2)

	p = doc.add_paragraph()
	_run(p, site, bold=True, size=30, color=_NAVY)
	p.paragraph_format.space_after = Pt(6)

	rule = doc.add_paragraph()
	_border(rule, color=_STEEL_HEX, size=12)
	rule.paragraph_format.space_after = Pt(14)

	meta = [
		("Facility", site),
		("ERP plant number", plant.erp_number or "\u2014"),
		("Source workbook", plant.workbook_filename or "\u2014"),
		("Generated", date.today().strftime("%B %d, %Y")),
		("Produced by", "ProcessArc \u2014 Phase 1"),
	]
	for label, value in meta:
		p = doc.add_paragraph()
		p.paragraph_format.space_after = Pt(3)
		_run(p, f"{label}   ", bold=True, size=10, color=_GREY, caps=True)
		_run(p, str(value), size=10.5, color=_INK)

	doc.add_paragraph()
	disclaimer = doc.add_paragraph()
	_shade(disclaimer, _NOTE_BG)
	disclaimer.paragraph_format.space_before = Pt(6)
	disclaimer.paragraph_format.space_after = Pt(6)
	disclaimer.paragraph_format.left_indent = Pt(8)
	disclaimer.paragraph_format.right_indent = Pt(8)
	_run(
		disclaimer,
		"Sequence text is reproduced from the customer source workbook. Wording "
		"is preserved verbatim except for clearly unambiguous spelling "
		"corrections, which are individually marked in the text and listed in "
		"Editorial Notes at the end of this document. Device names, numbers, and "
		"all operational content are never altered. This document reflects the "
		"source workbook as supplied.",
		italic=True, size=9, color=_GREY,
	)
	doc.add_page_break()


def _add_contents(doc: Document, headings: list[str]) -> None:
	h = doc.add_paragraph()
	_run(h, "Contents", bold=True, size=15, color=_NAVY)
	rule = doc.add_paragraph()
	_border(rule, color=_RULE, size=6)
	rule.paragraph_format.space_after = Pt(8)
	for heading in headings:
		p = doc.add_paragraph()
		p.paragraph_format.left_indent = Pt(12)
		p.paragraph_format.space_after = Pt(3)
		_run(p, heading, size=11, color=_INK)
	doc.add_page_break()


def _render_note(doc: Document, line: ClassifiedLine, indent: int = 20) -> None:
	p = doc.add_paragraph()
	_shade(p, _AMBER_BG)
	p.paragraph_format.left_indent = Pt(indent)
	p.paragraph_format.right_indent = Pt(8)
	p.paragraph_format.space_before = Pt(4)
	p.paragraph_format.space_after = Pt(4)
	_run(p, "NOTE   ", bold=True, size=8.5, color=_AMBER_LABEL, caps=True)
	_run(p, line.text, size=9.5, color=_AMBER_TEXT)
	_mark_correction(p, line)


def _render_conditional(doc: Document, line: ClassifiedLine) -> None:
	p = doc.add_paragraph()
	p.paragraph_format.left_indent = Pt(34)
	p.paragraph_format.space_after = Pt(2)
	_run(p, "IF  ", bold=True, size=8.5, color=_COND_LABEL)
	_run(p, line.text, italic=True, size=10)
	_mark_correction(p, line)


def _render_transition(doc: Document, line: ClassifiedLine) -> None:
	p = doc.add_paragraph()
	p.paragraph_format.left_indent = Pt(20)
	p.paragraph_format.space_before = Pt(3)
	p.paragraph_format.space_after = Pt(6)
	_run(p, "\u2192  ADVANCES WHEN   ", bold=True, size=8.5, color=_GREY, caps=True)
	text = line.text
	for prefix in ("Step advance on ", "Step advance ", "Step completes "):
		if text.lower().startswith(prefix.lower()):
			text = text[len(prefix):]
			break
	_run(p, text, italic=True, size=10, color=_TRANS_TEXT)


def _render_numbered_action(doc: Document, n: int, line: ClassifiedLine) -> None:
	p = doc.add_paragraph()
	p.paragraph_format.left_indent = Pt(20)
	p.paragraph_format.space_after = Pt(2)
	_run(p, f"{n}.  ", bold=True, size=10, color=_STEEL)
	_run(p, line.text, size=10.5)
	_mark_correction(p, line)


_CLASS_ABBREV: dict[DeviceClass, str] = {
	DeviceClass.VALVE: "V",
	DeviceClass.CONTROL_VALVE: "CV",
	DeviceClass.PUMP: "P",
	DeviceClass.VFD_PUMP: "VP",
	DeviceClass.TANK: "T",
}


def _devices_referenced_in_step(
	body: list[ClassifiedLine],
	devices_by_name: dict[str, DeviceRecord],
) -> list[DeviceRecord]:
	"""Find every Device Model device whose base_name appears (whole-word) in
	the prose of this step's body. Order preserved by first mention so the
	summary reads in the order the engineer would meet them.
	"""
	if not devices_by_name:
		return []
	body_text = " ".join(l.text for l in body)
	seen: list[DeviceRecord] = []
	taken: set[str] = set()
	# Iterate by name length descending so a longer name ("TankV1") matches
	# before a shorter substring would steal the word boundary ("V1").
	names_sorted = sorted(devices_by_name.keys(), key=len, reverse=True)
	# But we also want stable insertion order by first appearance — scan the
	# text once and check each candidate.
	for match in re.finditer(r"[A-Za-z][A-Za-z0-9]*", body_text):
		word = match.group(0)
		# Case-insensitive exact match against any device base_name.
		for name in names_sorted:
			if name.lower() == word.lower() and name not in taken:
				taken.add(name)
				seen.append(devices_by_name[name])
				break
	return seen


def _render_device_summary(doc: Document, devices: list[DeviceRecord]) -> None:
	"""Compact strip of device chips referenced in this step.

	Each chip is `BaseName (ClassAbbrev)` in a small grey font. Skipped if
	there are no matches (e.g. an introductory step that names no devices).
	"""
	if not devices:
		return
	p = doc.add_paragraph()
	p.paragraph_format.left_indent = Pt(20)
	p.paragraph_format.space_before = Pt(0)
	p.paragraph_format.space_after = Pt(4)
	_run(p, "DEVICES   ", bold=True, size=8, color=_GREY, caps=True)
	separator_color = RGBColor(0xC0, 0xC0, 0xC0)
	for i, d in enumerate(devices):
		if i > 0:
			_run(p, "   •   ", size=9, color=separator_color)
		_run(p, d.base_name, bold=True, size=9.5, color=_STEEL)
		abbrev = _CLASS_ABBREV.get(d.device_class, "")
		if abbrev:
			_run(p, f" ({abbrev})", italic=True, size=8.5, color=_GREY)


def _render_step(doc: Document, number: int, header: ClassifiedLine,
                 body: list[ClassifiedLine],
                 devices_by_name: dict[str, DeviceRecord]) -> None:
	bar = doc.add_paragraph()
	_shade(bar, _STEP_BG)
	bar.paragraph_format.space_before = Pt(10)
	bar.paragraph_format.space_after = Pt(4)
	bar.paragraph_format.left_indent = Pt(4)
	_run(bar, f"STEP {number}", bold=True, size=9, color=_STEEL, caps=True)
	_run(bar, "    ", size=9)
	_run(bar, header.text.rstrip(":"), bold=True, size=12, color=_NAVY)

	# Tier 2: per-step device summary strip (post-review Device Model names).
	devices_here = _devices_referenced_in_step(body, devices_by_name)
	_render_device_summary(doc, devices_here)

	action_n = 0
	for line in body:
		if line.kind in (LineType.ACTION, LineType.PROSE):
			action_n += 1
			_render_numbered_action(doc, action_n, line)
		elif line.kind == LineType.CONDITIONAL:
			_render_conditional(doc, line)
		elif line.kind == LineType.NOTE:
			_render_note(doc, line, indent=20)
		elif line.kind == LineType.TRANSITION:
			_render_transition(doc, line)
		# GROUP_HEADER / SECTION_TITLE inside a cylinder step are not expected;
		# if they appear, fall through silently rather than mis-render.


def _render_headerless(doc: Document, classified: list[ClassifiedLine]) -> None:
	"""Render a section with no step headers as one numbered procedure.

	Used for mix sheets, which the customer writes as a flat numbered list.
	Group headers ("For mixing Tanks 3 and 5:") render as subheadings;
	action numbering continues across them.
	"""
	n = 0
	for line in classified:
		if line.kind == LineType.SECTION_TITLE:
			continue
		if line.kind == LineType.GROUP_HEADER:
			p = doc.add_paragraph()
			p.paragraph_format.space_before = Pt(8)
			p.paragraph_format.space_after = Pt(4)
			p.paragraph_format.left_indent = Pt(4)
			_run(p, line.text, bold=True, size=11, color=_NAVY)
			_mark_correction(p, line)
		elif line.kind == LineType.NOTE:
			_render_note(doc, line, indent=20)
		elif line.kind == LineType.CONDITIONAL:
			_render_conditional(doc, line)
		elif line.kind == LineType.TRANSITION:
			_render_transition(doc, line)
		else:  # ACTION or PROSE
			n += 1
			_render_numbered_action(doc, n, line)


def _devices_for_sheet(model: Optional[DeviceModel], sheet: IngestedSheet) -> dict[str, DeviceRecord]:
	"""Build a base_name → DeviceRecord lookup for this section's system.

	Excluded devices are skipped (they were marked out at Review and shouldn't
	appear in the per-step summary). Returns an empty dict if there's no
	Device Model — Tier 2 then degrades to "no summary strip" instead of
	failing the render.
	"""
	if model is None:
		return {}

	if sheet.kind == SheetKind.CYLINDER_SEQUENCE and sheet.cylinder_number is not None:
		target_system, target_number = SystemKind.CYLINDERS, sheet.cylinder_number
	elif sheet.kind == SheetKind.MIX_SEQUENCE:
		# Mix sheets use a label, not a system number. Match by ANY mix
		# system in the model — usually plants have just one or two.
		return {
			d.base_name: d
			for d in model.devices
			if d.system == SystemKind.MIXING and d.review_status != ReviewStatus.EXCLUDED
		}
	else:
		return {}

	return {
		d.base_name: d
		for d in model.devices
		if d.system == target_system
		and d.system_number == target_number
		and d.review_status != ReviewStatus.EXCLUDED
	}


def _render_section(doc: Document, heading: str, sheet: IngestedSheet,
                     idle: bool, devices_by_name: dict[str, DeviceRecord]) -> list[ClassifiedLine]:
	h = doc.add_paragraph()
	h.paragraph_format.space_before = Pt(8)
	h.paragraph_format.space_after = Pt(2)
	_run(h, heading, bold=True, size=18, color=_NAVY)
	if idle:
		_run(h, "    IDLE \u2014 NOT COMMISSIONED", bold=True, size=9, color=_GREY, caps=True)

	classified = classify_lines(sheet.text_lines())
	steps = [c for c in classified if c.kind == LineType.STEP_HEADER]

	meta = doc.add_paragraph()
	meta.paragraph_format.space_after = Pt(2)
	_run(meta, f"Source sheet: {sheet.name}", italic=True, size=9, color=_GREY)
	if steps:
		_run(meta, f"      {len(steps)} steps", italic=True, size=9, color=_GREY)
	else:
		_run(meta, "      numbered procedure", italic=True, size=9, color=_GREY)

	rule = doc.add_paragraph()
	_border(rule, color=_RULE, size=6)
	rule.paragraph_format.space_after = Pt(8)

	# No step headers (mix sheets) -> render as a single numbered procedure.
	if not steps:
		_render_headerless(doc, classified)
		return classified

	# Group lines into steps. Anything before the first step header is preamble.
	preamble: list[ClassifiedLine] = []
	grouped: list[tuple[ClassifiedLine, list[ClassifiedLine]]] = []
	cur_header: Optional[ClassifiedLine] = None
	cur_body: list[ClassifiedLine] = []
	for line in classified:
		if line.kind == LineType.SECTION_TITLE:
			continue  # redundant restatement of the heading — drop
		if line.kind == LineType.STEP_HEADER:
			if cur_header is not None:
				grouped.append((cur_header, cur_body))
			cur_header = line
			cur_body = []
		elif cur_header is None:
			preamble.append(line)
		else:
			cur_body.append(line)
	if cur_header is not None:
		grouped.append((cur_header, cur_body))

	for line in preamble:
		if line.kind == LineType.NOTE:
			_render_note(doc, line, indent=4)
		else:
			p = doc.add_paragraph()
			p.paragraph_format.left_indent = Pt(4)
			p.paragraph_format.space_after = Pt(2)
			_run(p, line.text, italic=True, size=10, color=_GREY)
			_mark_correction(p, line)

	for i, (header, body) in enumerate(grouped, start=1):
		_render_step(doc, i, header, body, devices_by_name)

	return classified


def _add_editorial_notes(doc: Document,
                         corrections: list[tuple[str, str]]) -> None:
	if not corrections:
		return
	doc.add_page_break()
	h = doc.add_paragraph()
	_run(h, "Editorial Notes", bold=True, size=15, color=_NAVY)
	rule = doc.add_paragraph()
	_border(rule, color=_RULE, size=6)
	rule.paragraph_format.space_after = Pt(8)

	p = doc.add_paragraph()
	_run(
		p,
		"ProcessArc corrected the following unambiguous spelling errors from the "
		"customer source workbook. Each is marked with \u2731 in the text above. "
		"No device names, numbers, or operational content were changed.",
		italic=True, size=9.5, color=_GREY,
	)
	doc.add_paragraph()

	seen: set[tuple[str, str]] = set()
	for was, now in corrections:
		key = (was.lower(), now.lower())
		if key in seen:
			continue
		seen.add(key)
		p = doc.add_paragraph()
		p.paragraph_format.left_indent = Pt(16)
		p.paragraph_format.space_after = Pt(3)
		_run(p, "\u2731  ", size=9, color=_STEEL)
		_run(p, f'"{was}"', size=10, color=_GREY)
		_run(p, "  \u2192  ", size=10, color=_GREY)
		_run(p, f'"{now}"', size=10, bold=True, color=_INK)


# --- public interface (unchanged signature & contract) ----------------------
def export_sequence_doc(
	plant: PlantConfiguration,
	sequence_workbook_path: Optional[Path],
	output_dir: str | Path,
	device_model: Optional[DeviceModel] = None,
) -> Optional[Path]:
	"""Render the Treating Sequence .docx.

	`device_model` is optional — when supplied, each step's body is scanned
	for device-name references and a small summary strip is rendered at the
	top of the step (Tier 2). When omitted, steps render without the strip.

	Returns the output Path, or None if the source workbook is missing on disk
	or contains no cylinder/mix sequencing sheets — the caller handles that as
	a missing-input case rather than a hard failure.
	"""
	if sequence_workbook_path is None or not Path(sequence_workbook_path).exists():
		return None

	wb = ingest_workbook(sequence_workbook_path)
	cylinder_sheets = [s for s in wb.sheets if s.kind == SheetKind.CYLINDER_SEQUENCE]
	mix_sheets = [s for s in wb.sheets if s.kind == SheetKind.MIX_SEQUENCE]
	if not cylinder_sheets and not mix_sheets:
		return None

	# Treating (cylinder) sections first, then mix sections.
	ordered = [(s, _section_heading(s)) for s in cylinder_sheets]
	ordered += [(s, _section_heading(s)) for s in mix_sheets]

	doc = Document()
	_setup_styles(doc)
	for section in doc.sections:
		section.page_width = Inches(8.5)
		section.page_height = Inches(11)
		section.top_margin = Inches(1)
		section.bottom_margin = Inches(1)
		section.left_margin = Inches(1)
		section.right_margin = Inches(1)

	_add_cover(doc, plant)
	_add_contents(doc, [heading for _, heading in ordered])

	all_corrections: list[tuple[str, str]] = []
	for idx, (sheet, heading) in enumerate(ordered):
		idle = _cylinder_is_idle(plant, sheet)
		devices_by_name = _devices_for_sheet(device_model, sheet)
		classified = _render_section(doc, heading, sheet, idle, devices_by_name)
		for line in classified:
			all_corrections.extend(line.corrections)
		if idx != len(ordered) - 1:
			doc.add_page_break()

	_add_editorial_notes(doc, all_corrections)

	out_dir = Path(output_dir)
	out_dir.mkdir(parents=True, exist_ok=True)
	out_path = out_dir / f"{_safe_site_name(plant)}_TreatingSequence.docx"
	doc.save(out_path)
	return out_path
